import os as o
import re
import spacy as sp
import json as js
import pandas as pd
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup as bs
from pdfminer.high_level import extract_text
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from docx import Document
import mailparser
from stix2 import parse as parse_stix
from datetime import datetime
from collections import OrderedDict

# Load spaCy model
nlp = sp.load("en_core_web_sm")

filepath = input('Enter the file path: ')
ext = o.path.splitext(filepath)[1].lower()

# ========= TEXT EXTRACTION ==========
def extract_text_from_file(filepath):
    try:
        if ext == '.pdf':
            return extract_text(filepath)

        elif ext in ['.html', '.htm']:
            with open(filepath, 'r', encoding='utf-8') as f:
                global soup
                soup = bs(f, 'lxml')
                return soup.get_text(separator='\n')

        elif ext == '.txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()

        elif ext == '.docx':
            doc = Document(filepath)
            return '\n'.join([para.text for para in doc.paragraphs])

        elif ext == '.csv':
            with open(filepath, newline='', encoding='utf-8') as f:
                return '\n'.join(f.read().splitlines())

        elif ext == '.xlsx':
            df = pd.read_excel(filepath)
            return df.to_string(index=False)

        elif ext == '.json':
            with open(filepath, 'r', encoding='utf-8') as f:
                data = js.load(f)
                return js.dumps(data, indent=2)

        elif ext in ['.xml', '.ioc']:
            tree = ET.parse(filepath)
            root = tree.getroot()
            return ET.tostring(root, encoding='unicode')

        elif ext in ['.stix', '.stix2']:
            with open(filepath, 'r', encoding='utf-8') as f:
                stix_obj = parse_stix(f.read())
                return js.dumps(stix_obj.serialize(), indent=2)

        elif ext == '.eml':
            email = mailparser.parse_from_file(filepath)
            headers = '\n'.join(f"{k}: {v}" for k, v in email.headers.items())
            return email.body + '\n' + headers

        else:
            print(f"[!] Unsupported File Type: {ext}")
            return ''

    except Exception as e:
        print(f"[!] Error reading {filepath}: {e}")
        return ''

# ========= GET TEXT ==========
fp = extract_text_from_file(filepath)
if not fp:
    print("[!] No data extracted. Exiting...")
    exit()

# ========= TEXT CLEANUP ==========
fp = re.sub(r'\n+', ' ', fp)
fp = re.sub(r'\s{2,}', ' ', fp)
fp = re.sub(r'(?<=\b\w)\s(?=\w\b)', '', fp)

# ========= REGEX IOCs ==========
try:
    raw_ips = re.findall(r'\b\d{1,3}(?:\.|\[\.\])\d{1,3}(?:\.|\[\.\])\d{1,3}(?:\.|\[\.\])\d{1,3}\b', fp)
except Exception as e:
    print("[!] Error extracting IPv4 addresses:", e)
    raw_ips = []

extract1 = {
    'ipv4': list(set(ip.replace('[.]', '.') for ip in raw_ips)),
    'ipv6': re.findall(r'\b(?:[0-9a-fA-F]{1,4}:){7}[0-9a-fA-F]{1,4}\b', fp),
    'url': re.findall(r'https?://[^\s"\']+', fp) + re.findall(r'\b(?:www\.)[a-zA-Z0-9-]+\.[a-zA-Z]{2,6}\b', fp),
    'md5': re.findall(r'\b[a-fA-F0-9]{32}\b', fp),
    'sha1': re.findall(r'\b[a-fA-F0-9]{40}\b', fp),
    'sha256': re.findall(r'\b[a-fA-F0-9]{64}\b', fp),
    'attack_id': re.findall(r'T\d{4}(?:\.\d{3})?', fp)
}

# ========= NLP ENTITIES ==========
doc = nlp(fp)
seen_entities = set()
extract2 = []
for ent in doc.ents:
    key = (ent.text, ent.label_)
    if key not in seen_entities:
        seen_entities.add(key)
        extract2.append({
            "text": ent.text,
            "label": ent.label_,
            "start": ent.start_char,
            "end": ent.end_char
        })

threat_actors = list(OrderedDict.fromkeys(
    ent.text for ent in doc.ents if ent.label_ in ["ORG", "PERSON", "PRODUCT"]
))

# ========= HTML METADATA ==========
extract3 = {}
if ext in ['.html', '.htm']:
    try:
        extract3 = {
            'title': soup.title.text if soup.title else None,
            'urls': [a['href'] for a in soup.find_all('a', href=True)],
            'images': [img['src'] for img in soup.find_all('img', src=True)],
            'author': soup.find('meta', attrs={'name': 'author'})['content'] if soup.find('meta', attrs={'name': 'author'}) else None,
            'report_date': soup.find('meta', attrs={'name': 'date'})['content'] if soup.find('meta', attrs={'name': 'date'}) else None,
            'source': soup.find('meta', attrs={'name': 'source'})['content'] if soup.find('meta', attrs={'name': 'source'}) else None
        }
    except Exception as e:
        print("[!] Error extracting HTML metadata:", e)

# ========= PDF METADATA ==========
extract4 = {}
if ext == '.pdf':
    try:
        with open(filepath, 'rb') as file:
            parser = PDFParser(file)
            document = PDFDocument(parser)
            if hasattr(document, 'info') and document.info:
                metadata = document.info[0]
                for key, value in metadata.items():
                    try:
                        extract4[key.decode('utf-8')] = value.decode('utf-8')
                    except:
                        extract4[key.decode('utf-8')] = str(value)
    except Exception as e:
        print("[!] PDF metadata extraction failed:", e)

# ========= NORMALIZED VECTOR ==========
def normalize_output(extract1, extract2):
    return {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "hashes": extract1["md5"] + extract1["sha1"] + extract1["sha256"],
        "ips": extract1["ipv4"] + extract1["ipv6"],
        "domains": [u for u in extract1["url"] if not u.startswith("http")],
        "ttp_ids": extract1["attack_id"],
        "confidence": 90
    }

normalized_output = normalize_output(extract1, extract2)

# ========= FINAL OUTPUT ==========
final_output = {
    'regex_iocs': extract1,
    'nlp_entities': extract2,
    'threat_actors': threat_actors,
    'html_metadata': extract3,
    'pdf_metadata': extract4,
    'normalized_vector': normalized_output
}

out_name = o.path.splitext(o.path.basename(filepath))[0] + '_output.json'

try:
    with open(out_name, 'w', encoding='utf-8') as f:
        js.dump(final_output, f, indent=4)
    print(f"[✓] Output saved to {out_name}")
    #print(js.dumps(final_output, indent=2))
except Exception as e:
    print("[!] Failed to save JSON:", e)

